"""ingest_office.py — Word / Excel / CSV -> page-numbered text.

Sibling to ingest.py's PDF path. A mixed tender pack often carries the scored
questions, pricing schedule or compliance checklist as .docx/.xlsx/.csv rather
than PDF. These readers turn each into the same `IngestedDoc`/`Page` shape PDFs
already produce, with a human-readable locator (heading/paragraph, sheet/row,
CSV row) baked into the page text — so the existing chunker, extractor and
reconcile pass need no changes, and every requirement can still cite where it
came from. No `source_rect` is ever produced here (PDF-only, see pipeline.py);
callers leave it null for these documents.
"""

from __future__ import annotations

import csv
from pathlib import Path

from .ingest import IngestedDoc, Page, PDFIngestError


def _doc(filename: str, blocks: list[str]) -> IngestedDoc:
    text = "\n\n".join(b for b in blocks if b.strip())
    return IngestedDoc(filename=filename, pages=[Page(number=1, text=text)])


def ingest_docx(path: str | Path) -> IngestedDoc:
    """Read a Word doc's headings, paragraphs and tables in document order."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")
    try:
        import docx
    except ImportError as exc:
        raise PDFIngestError("python-docx is not installed; cannot read .docx files") from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise PDFIngestError(
            f"Could not parse {path.name} — the .docx file may be corrupt or unsupported."
        ) from exc

    blocks: list[str] = []
    para_num = 0
    current_heading: str | None = None
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        para_num += 1
        style = (para.style.name if para.style else "") or ""
        if style.lower().startswith("heading"):
            current_heading = text
            continue
        tag = f"[DOCX paragraph {para_num} | heading: {current_heading}]" if current_heading \
            else f"[DOCX paragraph {para_num}]"
        blocks.append(f"{tag}\n{text}")

    for t_idx, table in enumerate(document.tables, start=1):
        for r_idx, row in enumerate(table.rows, start=1):
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue
            row_text = " | ".join(cells)
            blocks.append(f"[DOCX table {t_idx} row {r_idx}]\n{row_text}")

    return _doc(path.name, blocks)


def ingest_xlsx(path: str | Path) -> IngestedDoc:
    """Read every non-empty row of every sheet, keeping sheet name + row number."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"XLSX not found: {path}")
    try:
        import openpyxl
    except ImportError as exc:
        raise PDFIngestError("openpyxl is not installed; cannot read .xlsx files") from exc

    try:
        workbook = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    except Exception as exc:
        raise PDFIngestError(
            f"Could not parse {path.name} — the .xlsx file may be corrupt or unsupported."
        ) from exc

    blocks: list[str] = []
    try:
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                values = [("" if c.value is None else str(c.value).strip()) for c in row]
                if not any(values):
                    continue
                row_num = row[0].row
                col_end = openpyxl.utils.get_column_letter(len(row))
                row_text = " | ".join(values)
                blocks.append(
                    f"[XLSX {sheet.title} row {row_num} | A{row_num}:{col_end}{row_num}]\n{row_text}"
                )
    finally:
        workbook.close()

    return _doc(path.name, blocks)


def ingest_csv(path: str | Path) -> IngestedDoc:
    """Read every non-empty CSV row, numbered."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"CSV not found: {path}")

    blocks: list[str] = []
    try:
        with path.open(newline="", encoding="utf-8-sig", errors="replace") as f:
            for row_num, row in enumerate(csv.reader(f), start=1):
                values = [cell.strip() for cell in row]
                if not any(values):
                    continue
                blocks.append(f"[CSV row {row_num}]\n{', '.join(values)}")
    except Exception as exc:
        raise PDFIngestError(
            f"Could not parse {path.name} — the .csv file may be corrupt or unsupported."
        ) from exc

    return _doc(path.name, blocks)
